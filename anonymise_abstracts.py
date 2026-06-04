"""Extract, anonymise, and export abstract text from supported document files."""

import argparse
import csv
import json
import re
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree

import fitz  # PyMuPDF
from docx import Document

# TODO: implement logging
Row = dict[str, str]

# Locations for saving
INPUT_FOLDER: str = "abstracts_raw"
OUTPUT_FOLDER: str = "abstracts_clean"
CSV_OUTPUT_FOLDER: str = "abstract_csv"
JSON_OUTPUT_FOLDER: str = "abstract_json"
# File for storage
OUTPUT_CSV: str = "anonymised_abstracts.csv"
OUTPUT_JSON: str = "anonymised_abstracts.json"
SUPPORTED_SUFFIXES: set[str] = {
    ".doc",
    ".docx",
    ".md",
    ".pdf",
    ".pptx",
    ".rst",
    ".text",
    ".txt",
}
TEXT_SUFFIXES: set[str] = {".md", ".rst", ".text", ".txt"}
CSV_FIELDNAMES: list[str] = [
    "file_name",
    "file_type",
    "run_id",
    "emails_removed_count",
    "orcids_removed_count",
    "phones_removed_count",
    "clean_text_file",
    "clean_text",
]

# Patterns for personal/contact identifiers removed from extracted text.
EMAIL_RE: str = (
    r"(?<![\w.+-])"
    r"[A-Za-z0-9._%+-]+@"
    r"(?:[A-Za-z0-9-]+\.)+?"
    r"[A-Za-z]{2,}"
    r"(?=$|[^\w.+-]|\.(?=\s|$|[A-Z][a-z]))"
)
ORCID_RE: str = r"\b(?:https?://orcid\.org/)?\d{4}-\d{4}-\d{4}-\d{3}[\dX]\b"
PHONE_RE: str = r"(\+?\d[\d\s().-]{7,}\d)"


def parse_args() -> argparse.Namespace:
    """Parse command line options for all-file or single-file processing."""
    parser = argparse.ArgumentParser(
        description="Anonymise abstract text from supported document files."
    )
    parser.add_argument(
        "file_name",
        nargs="?",
        help=(
            "Optional file name or path to process. If omitted, "
            "all supported files in abstracts_raw are processed."
        ),
        type=str,
    )
    return parser.parse_args()


def extract_text_from_pdf(path: Path) -> str:
    """Return the combined page text from a PDF file."""
    text: list[str] = []
    with fitz.open(path) as pdf:
        for page in pdf:
            text.append(page.get_text())
    return "\n".join(text)


def extract_text_from_docx(path: Path) -> str:
    """Return paragraph and table-cell text from a DOCX file."""
    doc = Document(path)
    text = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.extend(p.text for p in cell.paragraphs)
    return "\n".join(part for part in text if part.strip())


def extract_text_from_doc(path: Path) -> str:
    """Return text from a legacy DOC file using a local extractor."""
    for command in ("antiword", "catdoc"):
        if shutil.which(command):
            result = subprocess.run(
                [command, str(path)],
                check=True,
                capture_output=True,
                text=True,
            )
            return result.stdout

    if shutil.which("strings"):
        result = subprocess.run(
            ["strings", "-n", "5", str(path)],
            check=True,
            capture_output=True,
            text=True,
        )
        return result.stdout

    raise ValueError(
        "Legacy .doc extraction requires antiword, catdoc, or strings."
    )


def extract_text_from_pptx(path: Path) -> str:
    """Return text from PPTX slide XML."""
    text: list[str] = []
    with zipfile.ZipFile(path) as presentation:
        slide_names = sorted(
            name for name in presentation.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        for slide_name in slide_names:
            root = ElementTree.fromstring(presentation.read(slide_name))
            for node in root.iter():
                if node.tag.endswith("}t") and node.text:
                    text.append(node.text)
    return "\n".join(text)


def extract_text_from_text_file(path: Path) -> str:
    """Return text from a UTF-8 compatible text file."""
    return path.read_text(encoding="utf-8", errors="replace")


def extract_text(path: Path) -> str:
    """Return extracted text from a supported input file."""
    match path.suffix.lower():
        case ".pdf":
            return extract_text_from_pdf(path)
        case ".docx":
            return extract_text_from_docx(path)
        case ".doc":
            return extract_text_from_doc(path)
        case ".pptx":
            return extract_text_from_pptx(path)
        case suffix if suffix in TEXT_SUFFIXES:
            return extract_text_from_text_file(path)
        case _:
            raise ValueError(f"Unsupported file type: {path.suffix}")


def anonymise_text(text: str) -> str:
    """Replace supported personal identifiers with removal markers."""
    clean_text, _ = anonymise_text_with_counts(text)
    return clean_text


def anonymise_text_with_counts(text: str) -> tuple[str, Row]:
    """Replace supported personal identifiers and count each replacement."""
    text, emails_removed = re.subn(EMAIL_RE, "[EMAIL_REMOVED]", text)
    text, orcids_removed = re.subn(ORCID_RE, "[ORCID_REMOVED]", text)
    text, phones_removed = re.subn(PHONE_RE, "[PHONE_REMOVED]", text)
    return text, {
        "emails_removed_count": str(emails_removed),
        "orcids_removed_count": str(orcids_removed),
        "phones_removed_count": str(phones_removed),
    }


def process_file(path: Path) -> Row:
    """Extract and anonymise a supported input file."""
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    raw_text = extract_text(path)
    clean_text, removal_counts = anonymise_text_with_counts(raw_text)
    return {
        "file_name": path.name,
        "file_type": suffix,
        **removal_counts,
        "clean_text": clean_text,
    }


def clean_text_path(output_dir: Path, source_path: Path) -> Path:
    """Build a clean-text output path for a source file."""
    suffix_name = source_path.suffix.lower().lstrip(".")
    return output_dir / f"{source_path.stem}_{suffix_name}_clean.txt"


def supported_input_files(input_dir: Path) -> list[Path]:
    """List supported input files from the raw abstracts directory."""
    return sorted(
        path for path in input_dir.iterdir()
        if (
            path.is_file()
            and not path.name.startswith("~$")
            and path.suffix.lower() in SUPPORTED_SUFFIXES
        )
    )


def resolve_input_path(file_name: str, input_dir: Path) -> Path:
    """Resolve a CLI file argument to one concrete supported input path."""
    candidate = Path(file_name)
    if not candidate.is_absolute() and candidate.parent == Path("."):
        candidate = input_dir / candidate

    if candidate.exists():
        return candidate

    if candidate.suffix:
        raise FileNotFoundError(candidate)

    # Allow a unique stem such as "abstract_001" only when it is unambiguous.
    matches = [
        path for path in supported_input_files(input_dir)
        if path.stem == candidate.name
    ]
    if not matches:
        raise FileNotFoundError(candidate)
    if len(matches) > 1:
        match_names = ", ".join(path.name for path in matches)
        raise ValueError(
            f"Multiple input files match '{file_name}': {match_names}. "
            "Please include the extension."
        )
    return matches[0]


def read_csv_rows(csv_path: Path) -> list[Row]:
    """Read existing aggregate rows from a CSV file."""
    if not csv_path.exists():
        return []
    csv.field_size_limit(10_000_000)
    with csv_path.open(newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        return [
            {field: row.get(field) or "" for field in CSV_FIELDNAMES}
            for row in reader
        ]


def read_json_rows(json_path: Path) -> list[Row]:
    """Read existing aggregate rows from a JSON file."""
    if not json_path.exists():
        return []
    with json_path.open(encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, list):
        return []
    rows: list[Row] = []
    for row in data:
        if isinstance(row, dict):
            rows.append({
                field: str(row.get(field) or "") for field in CSV_FIELDNAMES
            })
    return rows


def existing_rows(
    csv_path: Path,
    json_path: Path,
) -> list[Row]:
    """Find the first available aggregate output to seed single-file runs."""
    for path, reader in (
        (csv_path, read_csv_rows),
        (json_path, read_json_rows),
    ):
        rows = reader(path)
        if rows:
            return rows
    return []


def upsert_rows(existing: list[Row], processed: list[Row]) -> list[Row]:
    """Replace or add processed rows while preserving other existing rows."""
    rows_by_file = {row["file_name"]: row for row in existing}
    for row in processed:
        rows_by_file[row["file_name"]] = row
    return sorted(rows_by_file.values(), key=lambda row: row["file_name"])


def write_csv(rows: list[Row], csv_path: Path) -> None:
    """Write aggregate rows to CSV."""
    csv_path.parent.mkdir(exist_ok=True)
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=CSV_FIELDNAMES)
        writer.writeheader()
        writer.writerows(rows)


def write_json(rows: list[Row], json_path: Path) -> None:
    """Write aggregate rows to JSON."""
    json_path.parent.mkdir(exist_ok=True)
    with json_path.open("w", encoding="utf-8") as f:
        json.dump(rows, f, ensure_ascii=False, indent=2)
        f.write("\n")


def get_current_time_data() -> str:
    """Return date-time text for output logging."""
    now = datetime.now()
    return datetime.strftime(now, '%d_%m_%Y__%H_%M')


def main() -> None:
    """Run the abstract anonymisation pipeline from the CLI."""
    args = parse_args()
    input_dir = Path(INPUT_FOLDER)
    output_dir = Path(OUTPUT_FOLDER)
    csv_path = Path(CSV_OUTPUT_FOLDER) / OUTPUT_CSV
    json_path = Path(JSON_OUTPUT_FOLDER) / OUTPUT_JSON

    if not input_dir.exists():
        raise SystemExit(f"Input folder not found: {input_dir}")

    output_dir.mkdir(exist_ok=True)

    paths: list[Path]
    if args.file_name:
        try:
            paths = [resolve_input_path(args.file_name, input_dir)]
        except (FileNotFoundError, ValueError) as exc:
            raise SystemExit(str(exc)) from exc
        if paths[0].suffix.lower() not in SUPPORTED_SUFFIXES:
            supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
            raise SystemExit(
                f"Unsupported file type: {paths[0].suffix}. "
                f"Supported types: {supported}"
            )
    else:
        paths = supported_input_files(input_dir)

    processed_at = get_current_time_data()
    run_id = f"run_{processed_at}"
    rows: list[Row] = []
    for path in paths:
        result = process_file(path)
        clean_file = clean_text_path(output_dir, path)
        clean_file.write_text(result["clean_text"], encoding="utf-8")
        rows.append({
            "file_name": result["file_name"],
            "file_type": result["file_type"],
            "run_id": run_id,
            "emails_removed_count": result["emails_removed_count"],
            "orcids_removed_count": result["orcids_removed_count"],
            "phones_removed_count": result["phones_removed_count"],
            "clean_text_file": str(clean_file),
            "clean_text": result["clean_text"],
        })

    output_rows: list[Row] = rows
    if args.file_name:
        # Keep previous aggregate rows so targeted runs do not discard work.
        output_rows = upsert_rows(
            existing_rows(csv_path, json_path),
            rows
        )

    write_csv(output_rows, csv_path)
    write_json(output_rows, json_path)

    print(f"Processed {len(rows)} files.")
    print(f"CSV created: {csv_path}")
    print(f"JSON created: {json_path}")
    print(f"Clean text files saved in: {OUTPUT_FOLDER}")


if __name__ == "__main__":
    main()
