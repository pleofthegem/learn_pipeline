import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

import fitz
import pytest
from docx import Document


@pytest.fixture
def sample_paragraphs() -> list[str]:
    return [
        "Title: Sample abstract",
        "Email: author@example.com",
        "ORCID: https://orcid.org/0000-0002-1825-0097",
        "Phone: +44 20 7946 0958",
    ]


@pytest.fixture
def sample_text(sample_paragraphs: list[str]) -> str:
    return "\n".join(sample_paragraphs)


@pytest.fixture
def email_before_adjacent_sentence_text() -> str:
    return "test@email.com.This was a test"


def write_sample_pdf(path: Path, paragraphs: list[str]) -> None:
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "\n".join(paragraphs), fontsize=12)
    pdf.save(path)
    pdf.close()


def write_sample_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def write_sample_pptx(path: Path, paragraphs: list[str]) -> None:
    slide_text = "".join(
        f"<a:p><a:r><a:t>{escape(paragraph)}</a:t></a:r></a:p>"
        for paragraph in paragraphs
    )
    slide_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">'
        f"<p:cSld><p:spTree>{slide_text}</p:spTree></p:cSld>"
        "</p:sld>"
    )
    with zipfile.ZipFile(path, "w") as presentation:
        presentation.writestr("ppt/slides/slide1.xml", slide_xml)


@pytest.fixture
def sample_pdf_path(tmp_path: Path, sample_paragraphs: list[str]) -> Path:
    path = tmp_path / "abstract_001.pdf"
    write_sample_pdf(path, sample_paragraphs)
    return path


@pytest.fixture
def sample_docx_path(tmp_path: Path, sample_paragraphs: list[str]) -> Path:
    path = tmp_path / "abstract_001.docx"
    write_sample_docx(path, sample_paragraphs)
    return path


@pytest.fixture
def sample_docx_table_path(tmp_path: Path, sample_paragraphs: list[str]) -> Path:
    path = tmp_path / "abstract_table.docx"
    document = Document()
    document.add_paragraph("Top-level paragraph")
    table = document.add_table(rows=1, cols=1)
    for paragraph in sample_paragraphs:
        table.cell(0, 0).add_paragraph(paragraph)
    document.save(path)
    return path


@pytest.fixture
def sample_pptx_path(tmp_path: Path, sample_paragraphs: list[str]) -> Path:
    path = tmp_path / "abstract_001.pptx"
    write_sample_pptx(path, sample_paragraphs)
    return path


@pytest.fixture
def sample_txt_path(tmp_path: Path, sample_text: str) -> Path:
    path = tmp_path / "abstract_001.txt"
    path.write_text(sample_text, encoding="utf-8")
    return path


@pytest.fixture
def sample_input_dir(tmp_path: Path, sample_paragraphs: list[str]) -> Path:
    input_dir = tmp_path / "abstracts_raw"
    input_dir.mkdir()

    (input_dir / "abstract_001.doc").write_text(
        "\n".join(sample_paragraphs),
        encoding="utf-8",
    )
    write_sample_docx(input_dir / "abstract_001.docx", sample_paragraphs)
    write_sample_pdf(input_dir / "abstract_001.pdf", sample_paragraphs)
    write_sample_pptx(input_dir / "abstract_001.pptx", sample_paragraphs)
    (input_dir / "abstract_001.txt").write_text(
        "\n".join(sample_paragraphs),
        encoding="utf-8",
    )
    write_sample_pdf(input_dir / "abstract_002.pdf", sample_paragraphs)
    (input_dir / "~$abstract_003.docx").write_text(
        "Temporary Office lock file",
        encoding="utf-8",
    )
    (input_dir / "notes.csv").write_text("Unsupported input", encoding="utf-8")

    return input_dir
